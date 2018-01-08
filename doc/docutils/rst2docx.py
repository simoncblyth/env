#!/usr/bin/env python
"""
rst2docx.py
============

Converts a limited subset of RST source into 
OpenXML docx (ie Word XML document).

Usage::

    rst2docx.py /tmp/report.rst /tmp/report.docx
    doctree.py /tmp/report.rst


Dependencies
-------------

* python-docx see docx-




See also
---------

*docxbuilder-* which does a similar thing with Sphinx integration 
but using an obsolete version of docx and with template document
complications  

*openxml-* *mono-* based use the the C# OpenXML API, impressions
of C# : Java verbosity of code, Microsoft boring documentation  


Document debugging
---------------------


asciicheck.py : non-ascii character check
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

::

    simon:dec2017-lon blyth$ asciicheck.py index.rst
    index.rst 290     In the next 10 minutes you[?][?][?]ll receive an email containing your booking details and Hotel voucher/s.

    index.rst 303     Dec 21 - 23, 2017 [?][?][?] 1 guest


doctree.py : rst debugging by examination of intermediate pseudo-xml doctree
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

* generally fix issues using literal blocks to avoid rst ctrl chars 

::

    simon:dec2017-lon blyth$ rst2docx.py index.rst
    2017-12-18 13:55:38,664 env.doc.docutils.rst2docx INFO     appending docxpath : /Users/blyth/workflow/notes/travel/2017/dec2017-lon/index.docx 
    2017-12-18 13:55:38,664 env.doc.docutils.rst2docx INFO     reading /Users/blyth/workflow/notes/travel/2017/dec2017-lon/index.rst 
    /Users/blyth/workflow/notes/travel/2017/dec2017-lon/index.rst:158: (WARNING/2) Inline strong start-string without end-string.
    /Users/blyth/workflow/notes/travel/2017/dec2017-lon/index.rst:158: (WARNING/2) Inline strong start-string without end-string.
    /Users/blyth/workflow/notes/travel/2017/dec2017-lon/index.rst:158: (WARNING/2) Inline strong start-string without end-string.
    2017-12-18 13:55:38,893 env.doc.docutils.rst2docx INFO     Writer pre walkabout
    NotImplementedError: env.doc.docutils.rst2docx.Translator visiting unknown node type: problematic
    Exiting due to error.  Use "--traceback" to diagnose.
    Please report errors to <docutils-users@lists.sf.net>.
    Include "--traceback" output, Docutils version (0.12 [release]),
    Python version (2.7.13), your OS type & version, and the
    command line used.
    simon:dec2017-lon blyth$ 




"""

import os, logging, argparse
log = logging.getLogger(__name__)


from docx import Document
from docx.shared import Inches

from docutils import writers
from docutils.core import Publisher
import docutils.nodes as nodes



from docx.oxml.shared import OxmlElement, qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text  import WD_BREAK
from docx.shared import Inches 
from docx.shared import Pt



from translator import BaseTranslator

default_usage = '%prog [options] [<source> [<destination>]]'
default_description = ('Reads from <source> (default is stdin) and writes to '
                       '<destination> (default is stdout).  See '
                       '<http://docutils.sf.net/docs/user/config.html> for '
                       'the full reference.')


class AddTOC(object):
    """
    Attempt to follow the hint from 

    * https://github.com/python-openxml/python-docx/issues/36

    The document shows up, but no TOC (in Pages as least) despite the
    following gobbledegook being incorporated::

          <w:r>
            <w:fldChar w:fldCharType="begin"/>
          </w:r>
          <w:r>
            <w:instrText xml:space="preserve">TOC \* MERGEFORMAT</w:instrText>
          </w:r>
          <w:r>
            <w:fldChar w:fldCharType="end"/>
          </w:r>

    * http://stackoverflow.com/questions/9762684/how-to-generate-table-of-contents-using-openxml-sdk-2-0

    Hmm, it seems the effort required to automatically AddTOC 
    are not worth the effort. Workaround: do it manually within the client, it would 
    be necessary in any case to adjust styles of headings etc..

    """
    def __init__(self, docx):
        self.para = docx.add_paragraph()
        self.add_element(self.element('w:fldChar', {"w:fldCharType":'begin'}))
        self.add_element(self.element('w:instrText', {'xml:space':'preserve', '_text':"TOC \* MERGEFORMAT"}))
        self.add_element(self.element('w:fldChar', {'w:fldCharType':'end'}))

    def add_element(self, elem):
        run = self.para.add_run()
        r_element = run._r
        r_element.append(elem) 

    def __repr__(self):
        p_element = self.para._p
        return str(p_element.xml)

    def element(self, name, attr):
        from docx.oxml.shared import OxmlElement, qn
        elem = OxmlElement(name)  
        text = attr.pop('_text', None)
        for k, v in attr.items():
            elem.set(qn(k), v) 
        pass 
        if not text is None:
           elem.text = text 
        pass
        return elem



class Writer(writers.Writer):
    supported = ('pprint', 'pformat', 'pseudoxml')
    config_section = 'pseudoxml writer'
    config_section_dependencies = ('writers',)
    output = None

    def __init__(self, uribase):
        writers.Writer.__init__(self) 

        self.translator_class = Translator
        self.uribase = uribase
        self.docx = Document()

        self.make_verbatim_style("VerbatimStyle", char=False)  
        self.make_verbatim_style("CharVerbatimStyle", char=True)  


    def make_verbatim_style(self, name, char=False):

        if char:
            sty = self.docx.styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
        else:
            sty = self.docx.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH) 
        pass

        font = sty.font 
        font.name = 'Courier New'
        font.size = Pt(8)       

        if not char:
            pfmt = sty.paragraph_format
            pfmt.left_indent = Inches(0.0)
            pfmt.right_indent = Inches(0.0)
        pass

        return sty


    def resolve(self, uri):
        path = os.path.join(self.uribase, uri)
        return os.path.abspath(path)
 
    def translate(self):
        log.info("Writer pre walkabout")
        previsit = self.translator_class(self.document, self.docx, self, mode="pre", previsit=None)
        self.document.walkabout(previsit)
        #previsit.toc()

        log.info("Writer create walkabout")
        visitor = self.translator_class(self.document, self.docx, self, mode="create", previsit=previsit)
        self.document.walkabout(visitor)

        log.info("docinfo %s " % repr(visitor.docinfo))

        self.output = visitor.astext()
        self.visitor = visitor

        #log.info("pformat %s " % self.document.pformat())
        #log.info("output  %s " % self.output)


    def supports(self, format):
        """This writer supports all format-specific elements."""
        return True

    def save(self, path):
        log.info("save to %s " % path)
        self.docx.save(path)

        base, ext = os.path.splitext(path)
        pxml_path = "%s%s" % (base, ".pxml")

        log.info("save pseudo-xml to %s " % pxml_path)

        with open(pxml_path, "w") as fp:
            fp.write(self.document.pformat())
        


        



def node_astext(node):
    """avoids RST source linebreaks influencing output formatting"""
    return node.astext().replace("\n", " ")  


class Text(list):
    def __init__(self, translator):
        self.translator = translator
        list.__init__(self)

    def append(self, txt, msg=None):
        list.append(self, txt)
        len = list.__len__(self)
        if self.translator.incaption:
            log.info("append %s %s %s " % (len,msg, txt))

    def pop(self, msg=None):
        txt = list.pop(self)
        len = list.__len__(self)

        if self.translator.incaption:
            for i in range(len):
                print "%4d : %s " % (i, list.__getitem__(self, i))
            pass
            log.info("pop %s %s %s " % (len,msg, txt))
        pass
        return txt
   



class Translator(BaseTranslator):

    verbatim_start = "{\\small\\begin{verbatim}"
    verbatim_end = "\\end{verbatim}}"

    texttt_start = "\\texttt{"
    texttt_end = "}"

    def __init__(self, document, docx, writer, mode="create", previsit=None ):
        BaseTranslator.__init__(self, document)
        self.docx = docx
        self.parax = None

        pre = False
        if mode == "pre":
            pre = True
        pass
        self.pre = pre
        # "pre" mode used for preparatory walkabout 
        # to obtain section titles for a TOC 
        self.previsit = previsit
       
        self.writer = writer
        self.level = 0 

        self.intitle = False
        self.ininfo = False
        self.inraw = False
        self.incaption = False
        self.inliteral = False

        self.title_count = 0 
        self.titles = [] 
        self.caption_count = 0 
        self.pretext = None

        self.text = Text(self)
        self.para_style = []
        self.char_style = []
        self.docinfo = {}

        self.date = None
        self.newline = False 


    def astext(self):
        return ""


    def visit_section(self, node):
        self.level += 1
    def depart_section(self, node):
        self.level -= 1

    def visit_title(self, node):
        self.intitle = True
        self.title_count += 1  
    def depart_title(self, node):
        self.intitle = False

    def visit_paragraph(self, node):
        if self.ininfo or self.pre:return
        self.start_paragraph( "", self.parastyle )

    def start_paragraph(self, txt, parastyle):
        parax = self.docx.add_paragraph(txt, self.parastyle)
        self.parax = parax

    def depart_paragraph(self, node):
        if self.ininfo or self.pre:return
        self.parax = None

    def visit_caption(self, node):
        if self.pre:return
        self.caption_count += 1  
        self.incaption = True

        self.pretext = "Figure %s:" % self.caption_count

        parax = self.docx.add_paragraph("", self.parastyle)
        pfx = parax.paragraph_format 
        pfx.left_indent = Inches(0.25)
        pfx.space_after = Inches(0.25)
        #log.info("caption pfx %s " % repr(dir(pfx)))

        self.parax = parax

    def depart_caption(self, node):
        if self.pre:return
        self.parax = None
        self.incaption = False

    def visit_block_quote(self, node):
        log.info("visit_block_quote\n%s\n" % node_astext(node))
        parax = self.docx.add_paragraph("", self.parastyle)
        pfx = parax.paragraph_format 
        pfx.left_indent = Inches(1.25)
        pfx.space_after = Inches(0.25)
        self.parax = parax
  
    def depart_block_quote(self, node):
        log.info("depart_block_quote")
        self.parax = None

    def visit_comment(self, node):
        log.info("visit_comment\n%s\n" % node_astext(node))
  
    def depart_comment(self, node):
        log.info("depart_comment")


    def visit_Text(self, node):
        """
        Most text nodes appears inside paragraph nodes, the exceptions are titles and captions.
        """
        if self.ininfo or self.inraw or self.inliteral:return

        txt = node_astext(node)

        if self.pretext is not None:
           txt = "%s %s" % (self.pretext, txt)
           self.pretext = None 
        pass

        if self.parax is not None:
            if not self.pre:
                self.parax.add_run(txt, self.charstyle)
            pass
        elif self.intitle == True:
            log.debug("visit_Text whilst intitle") 
            ## somewhat specifically add date to first title heading
            if self.title_count == 1 and not self.date is None:
                txt = "%s [%s]" % (txt, self.date)
            pass
            self.titles.append([self.level, txt])
            if not self.pre:
                self.docx.add_heading(txt, self.level)
        else:
            if not self.pre:
                self.report("visit_Text without parax or intitle", node) 
                self.text.append(txt, msg="visit_Text")
                assert 0
        pass


    def depart_Text(self, node):
        pass

    def visit_enumerated_list(self, node):
        # list counters are not resetting, workaround: convert source document to bulleted 
        self.para_style.append('List Number')   
    def depart_enumerated_list(self, node):
        self.para_style.pop()

    def visit_bullet_list(self, node):
        self.para_style.append('List Bullet')
    def depart_bullet_list(self, node):
        self.para_style.pop()

    def visit_raw(self, node):

        if self.pre:return
        fmt = node.attributes.get('format',None)
        self.inraw = True

        text = node.astext()
        #txt = node_astext(node)
        if text == "\\newline":
            #log.info("visit_raw newline spotted")
            if self.parax is not None and not self.pre:
                self.parax.add_run("\n", self.charstyle)
            pass
        elif text == "TOC":
            if self.previsit is not None:
                log.info("visit_raw TOC spotted : inserting tocx")
                self.tocx(self.previsit) 
            pass
        elif text == "BREAK_PAGE":
            run = self.parax.add_run()
            run.add_break(WD_BREAK.PAGE)
        elif text.startswith(self.verbatim_start) and text.endswith(self.verbatim_end):

            btxt = text[len(self.verbatim_start):-len(self.verbatim_end)]  
            log.info("found verbatim txt block, lines %d " % len(btxt.split("\n")) )
            print "\n".join(btxt.split("\n"))
            self.docx.add_paragraph(btxt, style="VerbatimStyle")

        elif text.startswith(self.texttt_start) and text.endswith(self.texttt_end):
            ttxt = text[len(self.texttt_start):-len(self.texttt_end)]  

            self.parax.add_run(ttxt, "CharVerbatimStyle")
            #self.docx.add_paragraph(ttxt, style="VerbatimStyle")
            log.info("found inline texttt \"%s\" " % ttxt )
        else:
            log.info("skipping visit_raw node with text \"%s\" " % text) 
        pass
        #self.report("visit_raw", node) 
    def depart_raw(self, node):
        if self.pre:return
        #self.report("depart_raw", node) 
        self.inraw = False

    def visit_literal(self, node):
        txt = node_astext(node)

        log.info("visit_literal [%s]" % txt)
        if txt == "\\n":
            log.info("literal newline spotted")

        self.report("visit_literal", node) 
        pass
    def depart_literal(self, node):
        self.report("depart_literal", node) 

    def visit_literal_block(self, node):
        """
        Although the source is cleaner with literal blocks rather than 
        repurposing the latex visit_raw am continuing to 
        use the latter as that means can use the standard Sphinx rst to latex
        machinery and from the same source produce the docx with this rst2docx.py 
        script
        """
        if self.pre:return
        text = node.astext()
        log.info("visit_literal_block lines [%d]" % len(text.split("\n")))
        self.inliteral = True
        self.docx.add_paragraph(text, style="VerbatimStyle")
        log.info("visit_literal_block [%s]" % text)
        
    def depart_literal_block(self, node):
        if self.pre:return
        self.inliteral = False
        log.info("depart_literal_block")

    def toc(self):
        for level, label in self.titles:
            indent = " " * level
            log.info( " %s %s %s " % (indent, level, label)) 

    def tocx_entry(self, level, label):
        indent = " " * level
        log.info( "tocx %s %s %s " % (indent, level, label)) 
        charstyle = "Heading %d Char" %  level  
        parax = self.docx.add_paragraph("", style=None)
        pfx = parax.paragraph_format 
        pfx.left_indent = Inches(0.2*level)
        pfx.space_after = Inches(0.05)
        parax.add_run( indent + label, charstyle)
 
    def tocx(self, previsit):
        self.tocx_entry(1, "Contents") 
        for level, label in previsit.titles:
            if level > 1:
                self.tocx_entry(level, label) 
            pass
 
    def report(self, msg, node):
        txt = node_astext(node)
        log.info("report %s %s %s " % (msg, len(self.text), txt)) 

    def _parastyle(self):
        """
        https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html

        Examples of paragraph styles:
 
           List Bullet 
           List Number

        """
        if len(self.para_style): 
            style = self.para_style[-1]
        else:
            style = None
        pass
        return style
    parastyle = property(_parastyle)


    def _charstyle(self):
        """
        https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html

        Examples of Character styles in default template

            Body Text Char
            Body Text 2 Char
            Body Text 3 Char
            Book Title
            Default Paragraph Font
            Emphasis
            Heading 1 Char
            Heading 2 Char
            Heading 3 Char
            Heading 4 Char
            Heading 5 Char
            Heading 6 Char
            Heading 7 Char
            Heading 8 Char
            Heading 9 Char
            Intense Emphasis
            Intense Quote Char
            Intense Reference
            Macro Text Char
            Quote Char
            Strong
            Subtitle Char
            Subtle Emphasis
            Subtle Reference
            Title Char

        """
        if len(self.char_style): 
            style = self.char_style[-1]
        else:
            style = None
        pass
        return style
    charstyle = property(_charstyle)


    def visit_emphasis(self, node):
        self.char_style.append("Emphasis")
    def depart_emphasis(self, node):
        self.char_style.pop()

    def visit_strong(self, node):
        self.char_style.append("Strong")
    def depart_strong(self, node):
        self.char_style.pop()


    def visit_figure(self, node):
        pass
    def depart_figure(self, node):
        pass

    def visit_image(self, node):
        path = self.writer.resolve(node.attributes['uri'])
        if not self.pre:
            self.docx.add_picture(path)
    def depart_image(self, node):
        pass

 
    def visit_docinfo(self, node):
        self.ininfo = True
    def depart_docinfo(self, node):
        self.ininfo = False
        if 'title' in self.docinfo and not self.pre:
            self.docx.add_heading(self.docinfo['title'], self.level)

    def visit_field(self, node):
        pass
    def depart_field(self, node):
        pass
    def visit_field_name(self, node):
        self.field_name = node_astext(node)
    def depart_field_name(self, node):
        pass
    def visit_field_body(self, node):
        if not self.field_name is None:
             field_body = node_astext(node)
             self.docinfo[self.field_name] = field_body 
             self.field_name = None
        pass 
    def depart_field_body(self, node):
        pass
 
    def visit_date(self, node):
        self.date = node_astext(node)
    def depart_date(self, node):
        pass
  


class Config(object):
    def __init__(self, doc):
        parser = argparse.ArgumentParser(doc)
        parser.add_argument( "paths", nargs='+',  help=""  )
        self.args = parser.parse_args()



def main():
    logging.basicConfig(level=logging.INFO, format="%(asctime)s %(name)s %(levelname)-8s %(message)s" )

    uribase = os.environ.get('URIBASE',"/Library/WebServer/Documents")

    config = Config(__doc__)
    paths = map(lambda _:os.path.abspath(_), config.args.paths)

    if len(paths) == 1:
        assert paths[0].endswith(".rst")
        rstpath = paths[0]
        assert os.path.isfile(rstpath)
        docxpath = rstpath.replace(".rst",".docx") 
        log.info("appending docxpath : %s " % docxpath )
        paths.append(docxpath)
    pass

    assert len(paths) == 2

    log.info("reading %s " % paths[0])

    odir = os.path.dirname(paths[1])
    if not os.path.exists(odir):
        os.makedirs(odir)




    reader=None
    parser=None
    writer=Writer(uribase=uribase)

    reader_name='standalone'
    parser_name='restructuredtext'
    writer_name= None

    settings = None
    settings_spec = None
    settings_overrides = None
    config_section = None
    enable_exit_status=True

    argv=[paths[0]]
    usage=default_usage
    description=default_description
 
    pub = Publisher(reader, parser, writer, settings=settings)
    pub.set_components(reader_name, parser_name, writer_name)

    output = pub.publish(
        argv, usage, description, settings_spec, settings_overrides,
        config_section=config_section, enable_exit_status=enable_exit_status)


    writer.save(paths[1])


    return output



if __name__ == '__main__':
    main()

